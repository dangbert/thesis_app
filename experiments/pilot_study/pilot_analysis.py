#!/usr/bin/env python3
"""Script for analyzing/plotting data from pilot.xlsx (see backend/dump_pilot.py)"""

import argparse
import os
import json
import sys
import glob
from typing import Union
import pandas as pd
import Levenshtein as lev
from matplotlib import pyplot as plt
import docx

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
EXPERIMENTS_DIR = os.path.realpath(os.path.join(SCRIPT_DIR, ".."))
sys.path.append(EXPERIMENTS_DIR)
import config  # noqa: E402

logger = config.get_logger(__name__)


FNAME = os.path.join(config.ROOT_DIR, "backend/pilot.xlsx")


def main():
    parser = argparse.ArgumentParser(
        description="Utils for doing analysis on pilot.xlsx data (see backend/dump_pilot.py)",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    parser.add_argument(
        "--levenshtein",
        "-l",
        action="store_true",
        help="Compute levenshtein edit distances, and plot errors",
    )
    parser.add_argument(
        "--input",
        "-i",
        type=str,
        default=FNAME,
        help="Input file (pilot.xlsx)",
    )
    parser.add_argument(
        "--dump-for-translate",
        "-dt",
        action="store_true",
        help="Create a docx file for translating text fields within pilot.xlsx",
    )
    args = parser.parse_args()

    new_fname = "pilot_lev_dist_norm.xlsx"
    if args.levenshtein:
        levenshtein(args.input, new_fname)
        plot_evals(new_fname) # could read either file here
    elif args.dump_for_translate:
        dump_for_translate(args.input)
    else:
        parser.print_help()
        exit(1)


def _get_dfs(fname: str):
    assert os.path.isfile(fname)
    a1_df = pd.read_excel(fname, sheet_name="P1", index_col=0)
    a2_df = pd.read_excel(fname, sheet_name="S1", index_col=0)
    return a1_df, a2_df


def levenshtein(fname: str, new_fname: str):
    # a1_df, a2_df = _get_dfs(fname)
    excel_file = pd.ExcelFile(fname)
    sheets = {sheet_name: excel_file.parse(sheet_name) for sheet_name in excel_file.sheet_names}
    a1_df, a2_df = sheets["P1"], sheets["S1"]

    logger.info(
        "computing normalized Levenshtein distance between ai_feedback and human_feedback."
    )

    def get_lev(row: pd.Series) -> Union[float, None]:
        if not isinstance(row["ai_feedback"], str) or not isinstance(row["human_feedback"], str):
            return None
        return lev.ratio(row["ai_feedback"], row["human_feedback"])

    # add column lev_dist_norm from columns 'ai_feedback' and 'human_feedback'
    a1_df["lev_dist_norm"] = a1_df.apply(get_lev, axis=1)
    a2_df["lev_dist_norm"] = a2_df.apply(get_lev, axis=1)

    # TODO: consider regenerating AI feedback for S1 and comparing as a baseline
    # to see how much the AI feedback in P1 biased the final feedback
    config.safe_append_sheet(a1_df, "P1", new_fname)
    config.safe_append_sheet(a2_df, "S1", new_fname)
    logger.info(f"wrote '{new_fname}'")

    # merge dataframes, keeping rows with lev_dist_norm defined
    lev_df = pd.concat([a1_df, a2_df], ignore_index=True)
    lev_df = lev_df.dropna(subset=["lev_dist_norm"])
    plt.figure()
    ax1 = plt.gca()
    plt.title("Edit Distance Between AI and Human Feedback", pad=15)
    plt.ylabel("Normalized Levenshtein Distance")
    plt.boxplot(lev_df["lev_dist_norm"], showmeans=True, meanprops=config.MEANPROPS)

    # Adjust figure size and padding
    plt.gcf().set_size_inches(4, 4)
    # add more padding on the right for labels
    plt.subplots_adjust(left=0.28, right=0.72, top=0.90, bottom=0.05)
    plt.ylim(0, 1)

    # Hide x-axis
    ax1.axes.get_xaxis().set_visible(False)

    # Create a secondary y-axis
    ax2 = ax1.twinx()
    ax2.set_ylim(ax1.get_ylim())  # Ensure the secondary y-axis aligns with the primary y-axis

    # Add labels at y=0 and y=1 on the secondary y-axis
    ax2.text(1.05, 0, "max changes", transform=ax2.get_yaxis_transform(), ha='left', va='center')
    ax2.text(1.05, 1, "no changes", transform=ax2.get_yaxis_transform(), ha='left', va='center')

    # Make sure secondary y-axis does not have any tick marks or labels
    ax2.tick_params(axis='y', which='both', left=False, right=False, labelright=False)

    plot_path = "lev_dist_norm.pdf"
    plt.savefig(plot_path)
    logger.info(f"wrote '{plot_path}'")

# these match EvalControls.tsx
ALL_PROBLEMS = {
  'Accuracy/relevance',
  'Feedback style/tone',
  'Structure',
  'Grammar',
  'Too wordy',
  'Too short',
  'Other',
}
# better name for plot
PROBLEM_NICKNAMES = { 'Feedback style/tone': 'Style/tone' }
LIKERT_STARS = {
  1: 'Very dissatisfied',
  2: 'Dissatisfied',
  3: 'Unsure',
  4: 'Satisfied',
  5: 'Very satisfied',
}

def plot_evals(fname: str):
    """Plot problem categorizations and Likert scale ratings."""
    df = pd.read_excel(fname, sheet_name="anon")

    counts = {PROBLEM_NICKNAMES.get(error, error): 0 for error in ALL_PROBLEMS}
    df_filtered = df[df["review_problems"].notnull()]
    for i, row in df_filtered.iterrows():
        errors = row["review_problems"].split(";")
        for error in errors:
            error = PROBLEM_NICKNAMES.get(error, error)
            counts[error] += 1
    total_problems = sum(counts.values())
    logger.info(f"{len(df_filtered)}/{len(df)} attempts tagged with review_problems")
    logger.info(f"{total_problems} problems reported in total (across tagged attempts)")

    ratings = df["review_rating"].dropna()
    logger.info(f"{len(ratings)} ratings reported in total")

    # sort keys by value largest to smallest
    counts = {k: v for k, v in sorted(counts.items(), key=lambda item: item[1], reverse=True)}
    
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(20, 8))  # Double the width
    # First plot: error classifications
    title_fontsize = 16
    axis_fontsize = 14
    ax1.set_title("LLM Feedback Problems Distribution", fontsize=title_fontsize)
    ax1.set_ylabel("Frequency", fontsize=axis_fontsize)
    ax1.bar(counts.keys(), counts.values())
    ax1.set_xticklabels(counts.keys(), rotation=90, fontsize=axis_fontsize)
    # add extra padding below for labels
    plt.subplots_adjust(left=0.05, right=0.95, top=0.95, bottom=0.27)

    # Second plot: bar plot of ratings
    ax2.set_title("LLM Feedback Ratings Distribution", fontsize=title_fontsize)
    ax2.set_ylabel("Frequency", fontsize=axis_fontsize)
    # ax2.set_xlabel("Rating", fontsize=axis_fontsize)
    ax2.hist(ratings, bins=range(1, 7), align='left', rwidth=0.8) # up to 6 for better visibility
    ax2.set_xticks(range(1, 6))
    ax2.set_xticklabels([f"{LIKERT_STARS[i]} - {i}" for i in range(1, 6)], rotation=90, fontsize=axis_fontsize)

    plot_path = "feedback_barplot.pdf"
    fig.tight_layout()
    plt.subplots_adjust(wspace=0.1)  # more space between plots
    plt.savefig(plot_path)
    logger.info(f"wrote '{plot_path}'")



def dump_for_translate(fname: str, out_path: str = "pilot_translate.docx"):
    a1_df, a2_df = _get_dfs(fname)
    document = docx.Document()
    for i, df in enumerate([a1_df, a2_df]):
        document.add_paragraph(f"##### assignment {i+1} #####")
        columns = ["attempt_goal", "attempt_plan", "ai_feedback", "human_feedback"]
        for i, row in df.iterrows():
            # given no index in xlsx, it took attempt_id (i) as the index
            row_id = row.get("attempt_id", i)
            text = f"\n=== {row_id} ==="
            for key in columns:
                # val = row[key] if row[key] and isinstance(row[key], str) else "" # handle nans (float defaults)
                val = row[key] if isinstance(row[key], str) else "nan"
                text += f"\n\n## {key}:\n{val}"
            document.add_paragraph(text)
            document.add_page_break()

    document.save(out_path)
    logger.info(f"wrote '{out_path}'")


if __name__ == "__main__":
    main()
