import os
import sys
import re
import docx
from datasets import Dataset

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
EXPERIMENTS_DIR = os.path.realpath(os.path.join(SCRIPT_DIR, "../.."))
sys.path.append(EXPERIMENTS_DIR)
import config  # noqa: E402

logger = config.get_logger(__name__, level="INFO")

EXPECTED_KEYS = ["input", "instruction", "output", "orig_index"]
# PATTERN_EXP = r"<===@(\d+)===>\s*<===0===>(.*?)</===0===>\s*<===1===>(.*?)</===1===>\s*<===2===>(.*?)</===2===>\s*</===@\1===>"
# after translation, the number of equals signs can occasionally drop from 3 -> 2 as a side effect
EQQ = r"[=]{2,}"  # 2 or more equals signs
PATTERN_EXP = rf"<{EQQ}@(\d+){EQQ}>\s*<{EQQ}0{EQQ}>(.*?)</{EQQ}0{EQQ}>\s*<{EQQ}1{EQQ}>(.*?)</{EQQ}1{EQQ}>\s*<{EQQ}2{EQQ}>(.*?)</{EQQ}2{EQQ}>\s*</{EQQ}@\1{EQQ}>"


def serialize_sample(data: dict) -> str:
    """
    Convert a dictionary into a custom string format as specified.
    we use the index of the keys because number should stay the same after translation, whereas "instruction" -> "instructie" etc with some variation.
    """
    keys = sorted(list(data.keys()))
    assert set(keys) == set(EXPECTED_KEYS)
    # remove "orig_index"
    keys.remove("orig_index")  # e.g. ["input", "instruction", "output"]
    serialized_str = f"<===@{data['orig_index']}===>\n"
    for key in keys:
        idx = keys.index(key)
        serialized_str += f"<==={idx}===>{data[key]}</==={idx}===>\n"
    serialized_str += f"</===@{data['orig_index']}===>"
    return serialized_str


def deserialize_sample(data_str: str) -> dict:
    """
    Convert the custom string format back into a dictionary.
    """
    import re

    match = re.match(PATTERN_EXP, data_str, re.DOTALL)
    if match:
        orig_index, input_str, instruction, output = match.groups()
        return {
            "orig_index": int(orig_index),
            "output": output.strip(),
            "input": input_str.strip(),
            "instruction": instruction.strip(),
        }
    raise ValueError("String does not match the expected format")


def serialize_dataset(dataset, fname: str, assert_sanity: bool = False):
    """Serialize dataset to a text file or a Word (.docx) document."""
    bad_count = 0
    char_count = 0

    def process_item(item: dict):
        nonlocal bad_count
        nonlocal char_count
        cereal = serialize_sample(item)
        if not deserialize_sample(cereal) == item:
            bad_count += 1
            if assert_sanity:
                assert deserialize_sample(cereal) == item
        char_count += len(cereal)
        return cereal

    if fname.endswith(".docx"):
        document = docx.Document()
        for i, item in enumerate(dataset):
            cereal = process_item(item)
            document.add_paragraph(cereal)
            # if i % 10 == 0:
            #     document.add_page_break()
            #     document.add_paragraph("")
        document.save(fname)

    else:
        with open(fname, "w") as f:
            for item in dataset:
                cereal = process_item(item)
                f.write(cereal + "\n")
                char_count += 1  # for newline

    if bad_count > 0:
        logger.warning(
            f"Failed to perfectly encode: {bad_count}/{len(dataset)} samples"
        )
    filesize = _get_filesize_mb(fname)
    logger.info(
        f"Serialized dataset to '{fname}' ({char_count:,} total chars), {filesize:.3f} MB"
    )


def _get_filesize_mb(filepath: str) -> float:
    """
    Get the filesize of a file in megabytes.
    """
    filesize_bytes = os.path.getsize(filepath)
    filesize_mb = filesize_bytes / 1024**2
    return filesize_mb


def deserialize_dataset(fname: str) -> Dataset:
    """
    Deserialize dataset from a text or Word (.docx) file using pattern matching.
    Afterwards calling this function you can write the dataset to a .jsonl file with
    dataset.to_json("/tmp/tmp.jsonl")
    """
    dlist = []
    if fname.lower().endswith(".docx"):
        logger.info(f"loading dataset from a Word document: '{fname}'")
        document = docx.Document(fname)
        data_str = "\n".join([p.text for p in document.paragraphs])
    else:
        logger.info(f"loading dataset from a text file: '{fname}'")
        with open(fname, "r") as f:
            data_str = f.read()

    pattern = re.compile(PATTERN_EXP, re.DOTALL)
    matches = pattern.finditer(data_str)
    for match in matches:
        serialized_sample = match.group(0)  # Extract the entire matched sample block
        deserialized_dict = deserialize_sample(serialized_sample)
        dlist.append(deserialized_dict)

    dataset = Dataset.from_list(dlist)
    logger.info(f"Deserialized dataset with {len(dataset)} rows")
    return dataset
